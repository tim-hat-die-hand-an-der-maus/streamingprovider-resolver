import dataclasses
import json
import os
import socket
import urllib.parse
from abc import abstractmethod, ABC
from collections import defaultdict
from typing import Dict
from typing import List, Optional

import bs4
import requests
import urllib3
import uvicorn
from bs4 import Tag
# noinspection PyPackageRequirements
from docx import Document
# noinspection PyPackageRequirements
from docx.opc.exceptions import PackageNotFoundError
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from thefuzz import fuzz

LOG_DIRECTORY = "/var/log/werstreamtes" if os.name != "nt" else os.path.join(os.getenv("APPDATA"), "werstreamtes")
LOG_FILENAME = "log.docx"
BASE_URL = os.getenv("BASE_URL") or "https://werstreamt.es"
SEARCH_PATH = os.getenv("SEARCH_PATH") or "/suggestTitle?term="


def log(message: str):
    print(message)
    path = os.path.join(LOG_DIRECTORY, LOG_FILENAME)
    try:
        document = Document(path)
    except PackageNotFoundError:
        document = Document()

    document.add_paragraph(message)
    document.save(path)


class SearchRequest(BaseModel):
    werstreamtesLink: str


class TitleSearchRequest(BaseModel):
    title: str
    year: Optional[int] = None


app = FastAPI()


def search_vodster_by_title(title_query: str):
    url = "/".join([BASE_URL, SEARCH_PATH, title_query])
    return requests.get(url)


@dataclasses.dataclass
class SearchItem:
    title: str
    id: Optional[int] = None
    year: Optional[int] = None
    type: Optional[str] = None

    @classmethod
    def from_json_item(cls, key: str, js: Dict) -> 'SearchItem':
        _id = int(key[3:])
        title = js['value']
        label = js['label']
        soup = bs4.BeautifulSoup(label, "lxml")
        span = soup.find("span")
        res = span.text.strip().split(", ")
        if len(res) != 2:
            _type = res[0]
            year = None
        else:
            _type, year = res
            year = int(year)
        _type = _type.split("\n")[-1].strip()
        return cls(title, _id, year, _type)

    def to_json(self):
        return {
            "id": self.id,
            "title": self.title,
            "year": self.year,
            "type": self.type,
        }


@dataclasses.dataclass
class StreamProvider:
    id: Optional[str]
    name: str


@dataclasses.dataclass
class SearchProvider(ABC):
    name: str

    @abstractmethod
    def search(self, request: SearchRequest, **kwargs) -> Optional[List[SearchItem]]:
        pass


@dataclasses.dataclass
class Provider(ABC):
    name: str

    @abstractmethod
    def get_streaming_providers(self, info: str, **kwargs):
        pass


@dataclasses.dataclass
class PlexResolverMovie:
    title: str
    year: int

    @classmethod
    def from_json(cls, _json: Dict) -> "PlexResolverMovie":
        return cls(
            _json["title"],
            _json["year"]
        )


@dataclasses.dataclass
class PlexResolverResponseItem:
    name: str
    movies: List[PlexResolverMovie]
    error: Optional[str]  # is this correct?

    @classmethod
    def from_json(cls, _json: Dict) -> "PlexResolverResponseItem":
        return cls(
            _json["name"],
            [PlexResolverMovie.from_json(movie) for movie in _json.get("movies", [])],
            _json["error"]
        )


@dataclasses.dataclass
class Plex(Provider, SearchProvider):
    def __init__(self):
        super().__init__("plex")
        self.url = os.getenv("PLEX_RESOLVER_URL") or "http://plex-resolver:8080/movies"

    def get_movies(self, url: str = None) -> Optional[List[PlexResolverResponseItem]]:
        if not url:
            url = self.url

        try:
            result = requests.get(url)
        except (requests.exceptions.ConnectionError, socket.gaierror, urllib3.exceptions.MaxRetryError) as e:
            log(f"Failed to retrieve {url} due to {e}")
            return None

        return [PlexResolverResponseItem.from_json(j) for j in result.json().get("data", [])]

    def get_streaming_providers(self, info: str, **kwargs):
        pass

    def search(self, request: TitleSearchRequest, **kwargs) -> Optional[Dict[str, List[PlexResolverMovie]]]:
        results = defaultdict(list)
        response = self.get_movies()
        if not response:
            return None

        for response in response:
            movies = response.movies

            for movie in movies:
                if fuzz.token_set_ratio(request.title, movie) > 80:
                    item = SearchItem(title=movie.title)
                    if request.year:
                        if request.year == movie.year:
                            item.year = movie.year
                            results[response.name].append(item)
                    else:
                        results[response.name].append(item)

        return results


@dataclasses.dataclass
class WerStreamtEs(Provider, SearchProvider):
    def __init__(self):
        super().__init__("werstreamt.es")

    def get_streaming_providers(self, info: str, **kwargs) -> Optional[List[StreamProvider]]:
        try:
            result = requests.get(info)
            body = result.text
        except (requests.exceptions.ConnectionError, socket.gaierror, urllib3.exceptions.MaxRetryError):
            log("Failed to retrieve {} due to")
            return None

        soup = bs4.BeautifulSoup(body, "lxml")
        provider_elements: List[Tag] = soup.find_all(attrs={"class": "provider"})[1:]

        providers = []
        for element in provider_elements:
            name_element = element.find("a", attrs={"class": "left"})
            # value for sky is `Sky Go\nsky` for example
            name = name_element.text.strip().splitlines()[0]

            options = json.loads(element.get("data-options"))
            stream_provider_id = options.get("StreamProviderID")
            providers.append(StreamProvider(stream_provider_id, name))

        return providers

    def search(self, request: TitleSearchRequest, **kwargs) -> Optional[List[Dict]]:
        title = urllib.parse.quote(request.title)
        url = "https://www.werstreamt.es/suche/suggestTitle?term=" + title

        req = requests.get(url, headers={"Accept": "application/json"})
        if req.ok:
            js = req.json()
            results = [SearchItem.from_json_item(key, value).to_json() for key, value in js.items() if
                       key.startswith("id-")]
            if request.year is not None:
                results = [item for item in results if item["year"] == request.year]

            return results

        return None


@app.post("/search")
def movie_by_title(req: TitleSearchRequest):
    results = {}
    providers = [WerStreamtEs(), Plex()]

    for provider in providers:
        if result := provider.search(req):
            if isinstance(result, dict):
                for name, movies in result.items():
                    results[f"{provider.name}-{name}"] = movies
            else:
                results[provider.name] = result

    if not results:
        raise HTTPException(status_code=404, detail="Title not found")

    return results


@app.post("/")
def movie_by_link(req: SearchRequest):
    results = {}
    providers = [WerStreamtEs()]

    for provider in providers:
        results[provider.name] = provider.get_streaming_providers(req.werstreamtesLink)

    return results


if __name__ == "__main__":
    LOG_DIRECTORY = "werstreamtes"
    if not os.path.exists(LOG_DIRECTORY):
        try:
            os.makedirs(LOG_DIRECTORY)
        except PermissionError:
            print(f"No permission to log to {LOG_DIRECTORY}, logging to ./werstreamtes")
            LOG_DIRECTORY = "./werstreamtes"
            os.makedirs(LOG_DIRECTORY, exist_ok=True)

    log("Starting")
    # noinspection PyTypeChecker
    uvicorn.run(app)
